# app.py by Yuming 2024
import gradio as gr
import re
import pandas as pd
from docx import Document

# Extract unique citation pairs (Name, Year) + full text from the thesis document
def extract_citations_from_docx(doc_path):
    document = Document(doc_path)
    full_text = []

    # Extract text from paragraphs
    for para in document.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # Avoid duplication
                if cell.text not in full_text:
                    full_text.append(cell.text)
    text = ' '.join(full_text)

    # Regular expressions for various citation formats (according to APA7)
    patterns = [
        r'\b([A-Z][a-z]+) et al\. \((\d{4})\)',  # Narrative citation, More than three authors: e.g., Munikar et al. (2019)
        r'\b([A-Z][a-z]+) \((\d{4})\)',           # Narrative citation, Single / Two Author Citation e.g., Park (2023), Li and Hu (2024)
        r'\(([A-Z][a-z]+) et al\., (\d{4})\)',    # Parenthetical citation, More than three authors:  e.g., (Liu et al., 2023)
        r'\(([A-Z][a-z]+) & ([A-Z][a-z]+), (\d{4})\)',  # Parenthetical citation, Two Authors Citation e.g., (Batrinca & Treleaven, 2015)
        r'\(([A-Z][a-z]+), (\d{4})\)',            #  Parenthetical citation, Single Author Citation: e.g., (Weber, 1987)
        r"\b\(([A-Z][a-z]+( et al.)?,\s\d{4};?\s?)+\)\b", # Multiple Citations, e.g., (Weber, 1987; Orlikowski & Iacono, 2001; Li et al., 2024)
        r"\b\([A-Z][a-z]+,\s\d{4}[a-z]?(,\s\d{4}[a-z]?)?\)\b" # Same Author Multiple Works e.g., (Smith, 2020a, 2020b)
    ]

    citations = set()
    for pattern in patterns:
        matches = re.findall(pattern, text)
        for match in matches:
            if len(match) == 2:
                citations.add((match[0], match[1]))
            elif len(match) > 2:
                if ';' in match[0]:
                    sub_matches = re.findall(r'([A-Z][a-z]+)[^,;]*?, (\d{4})', match[0])
                    for sub_match in sub_matches:
                        citations.add((sub_match[0], sub_match[1]))
                else:
                    citations.add((match[1], match[2]))
            elif ';' in match:
                sub_matches = re.findall(r'([A-Z][a-z]+)[^,;]*?, (\d{4})', match)
                for sub_match in sub_matches:
                    citations.add((sub_match[0], sub_match[1]))

    return citations, text

# Extract references from the reference list in the Excel file
def references_from_excel(excel_path):
    df = pd.read_excel(excel_path, header=None)
    references = [str(cell) for cell in df[0] if pd.notna(cell)]
    return references

# Extract references in the format of (Name, Year)
def extract_references_from_excel(excel_path):
    df = pd.read_excel(excel_path, header=None)
    references = []

    for cell in df[0][1:]:
        if pd.isna(cell):
            continue
        cell_text = str(cell)
        name_match = re.match(r'([A-Z][a-z]+)', cell_text)
        year_match = re.search(r'\((\d{4})\)', cell_text)
        if name_match and year_match:
            references.append((name_match.group(1), year_match.group(1)))

    return references

# Check if thesis citations are present in the reference list
def check_citations_in_references(citations, references):
    citations_not_in_references = []
    for name, year in citations:
        found = False
        for cell in references:
            if name in cell and year in cell:
                found = True
                break
        if not found:
            citations_not_in_references.append((name, year))

    return citations_not_in_references

# Check if references are present in the thesis text
def check_references_in_citations(references, text):
    references_not_in_text = []

    for name, year in references:
        same_bracket_pattern = fr'\({name}[^)]*?{year}[^)]*?\)'
        same_bracket_pattern2 = fr'\(([^\)]*{name}[^,;]*?,\s*{year}[^)]*)\)'
        name_followed_by_year_pattern = fr'\b{name}\b(?:[^\(\)]*?\(\s*{year}\s*\))'

        if not (re.search(same_bracket_pattern, text) or re.search(same_bracket_pattern2, text) or re.search(name_followed_by_year_pattern, text)):
            references_not_in_text.append((name, year))

    return references_not_in_text

# Gradio interface function
def analyze_files(doc_file, excel_file):
    citations, text = extract_citations_from_docx(doc_file.name)
    references_all = references_from_excel(excel_file.name)
    references = extract_references_from_excel(excel_file.name)

    citations_not_in_references = check_citations_in_references(citations, references_all)
    references_not_in_text = check_references_in_citations(references, text)

    result = ""
    if citations_not_in_references:
        result += "Citations in the thesis not found in the reference list:\n"
        result += '\n'.join([f"{name} ({year})" for name, year in citations_not_in_references]) + "\n"
    else:
        result += "All citations in the thesis were found in the reference list.\n"

    if references_not_in_text:
        result += "\nReferences in the list not found in the thesis:\n"
        result += '\n'.join([f"{name} ({year})" for name, year in references_not_in_text]) + "\n"
    else:
        result += "All references in the list were found in the thesis.\n"

    return result

# Gradio UI setup
iface = gr.Interface(
    fn=analyze_files,
    inputs=[gr.File(label="Upload .docx file"), gr.File(label="Upload .xlsx file")],
    outputs="text",
    title="Citation and Reference Checker",
    description="Please upload your thesis (without Reference list) as .docx file, and an .xlsx reference list to check for citation and reference consistency."
)

iface.launch(share=True)
